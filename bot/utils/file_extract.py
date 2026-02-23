"""Извлечение текста из файлов."""
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_bytes(data: bytes, filename: str = "") -> Optional[str]:
    """
    Извлечь текст из файла по расширению.
    Поддерживает: .txt, .md, .pdf, .docx, .xlsx, .xls
    """
    ext = (filename or "").lower().split(".")[-1] if "." in (filename or "") else ""

    if ext in ("txt", "md", "text"):
        try:
            return data.decode("utf-8", errors="replace")
        except Exception as e:
            logger.warning("Failed to decode text file: %s", e)
            return None

    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n\n".join(parts).strip() or None
        except Exception as e:
            logger.warning("Failed to extract PDF text: %s", e)
            return None

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n\n".join(parts).strip() or None
        except Exception as e:
            logger.warning("Failed to extract DOCX text: %s", e)
            return None

    if ext == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    parts.append(" | ".join(str(c) if c is not None else "" for c in row))
            wb.close()
            return "\n\n".join(parts).strip() or None
        except Exception as e:
            logger.warning("Failed to extract XLSX text: %s", e)
            return None

    if ext == "xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=data)
            parts = []
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    parts.append(" | ".join(str(sheet.cell_value(row_idx, col)) for col in range(sheet.ncols)))
            return "\n\n".join(parts).strip() or None
        except Exception as e:
            logger.warning("Failed to extract XLS text: %s", e)
            return None

    return None
